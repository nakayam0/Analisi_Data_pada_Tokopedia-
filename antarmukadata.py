import streamlit as st
import pandas as pd
import plotly.express as px
import io
import textwrap
import time
from docx import Document
from docx.shared import Inches
import tempfile
import plotly.io as pio

# Opsional: set default format agar tidak perlu menulis format setiap kali
pio.kaleido.scope.default_format = "png"


# --- KONFIGURASI ---
st.set_page_config(page_title="SCRAPING PRODUK TOKOPEDIA", layout="wide")
st.markdown(
    """
    <style>
        body { background-color: #1E1E1E; color: white; }
        .stDataFrame { background-color: #1E1E1E; }
    </style>
    """,
    unsafe_allow_html=True
)

st.title("DASHBOARD PRODUK")

# --- PLACEHOLDER LABELS UNTUK SELECTBOX ---
ORDER_PLACEHOLDER = "— Pilih urutan —"
VALUE_PLACEHOLDER = "— Pilih kategori —"

# --- INISIALISASI SESSION STATE DEFAULT ---
def init_state():
    st.session_state.setdefault("filter_produk", "")
    st.session_state.setdefault("filter_perangkat", [])
    st.session_state.setdefault("filter_jenis", [])
    st.session_state.setdefault("order_option_label", ORDER_PLACEHOLDER)
    st.session_state.setdefault("value_option_label", VALUE_PLACEHOLDER)
    st.session_state.setdefault("default_col_option", "TERJUAL")
    st.session_state.setdefault("data_option", "Running Text")
init_state()

# --- FUNGSI RESET ---
def reset_filters():
    st.session_state["filter_produk"] = ""
    st.session_state["filter_perangkat"] = []
    st.session_state["filter_jenis"] = []
    st.session_state["order_option_label"] = ORDER_PLACEHOLDER
    st.session_state["value_option_label"] = VALUE_PLACEHOLDER
    st.session_state["default_col_option"] = "TERJUAL"
    st.session_state["reset_success"] = True

# --- 0. PILIH SUMBER DATA + RESET (SEJAJAR) ---
colA, colB = st.columns([4, 1])
with colA:
    data_option = st.selectbox(
        "📂 Pilih Sumber Data",
        ["Running Text", "Mesin Antrian", "Jadwal Sholat", "Videotron"],
        key="data_option"
    )
with colB:
    st.write("")
    st.button("🔄 Reset Filter", on_click=reset_filters)

# --- POPUP RESET ---
if st.session_state.get("reset_success", False):
    st.toast("✅ Reset Berhasil!", icon="🔄")
    time.sleep(2)
    st.session_state["reset_success"] = False
    st.rerun()

# Mapping pilihan ke file Excel
file_mapping = {
    "Running Text": "RUNNING TEXT_TOKOPEDIA.xlsx",
    "Mesin Antrian": "MESIN ANTRIAN.xlsx",
    "Jadwal Sholat": "SCRAPPING JWS JADI.xlsx",
    "Videotron": "SCRAPING VIDEOTRON TOKOPEDIA JADI.xlsx"
}

# --- BACA DATA SESUAI PILIHAN ---
file_path = file_mapping[data_option]
df = pd.read_excel(file_path)

# --- 1. DATA AWAL ---
st.subheader("Data Awal")
st.dataframe(df, use_container_width=True)

# --- 2. FILTER PERTAMA ---
col1, col2, col3 = st.columns(3)

with col1:
    filter_produk = st.text_input(
        "🔍 Filter Produk",
        placeholder="Masukkan produk",
        key="filter_produk"
    )

with col2:
    filter_perangkat = None
    if data_option == "Running Text":
        filter_perangkat = st.multiselect(
            "🔍 Filter Perangkat",
            ["running text", "modul", "strobo", "power supply"],
            default=st.session_state["filter_perangkat"],
            placeholder="Pilih opsi",
            key="filter_perangkat"
        )
    elif data_option == "Videotron":
        filter_perangkat = st.multiselect(
            "🔍 Filter Perangkat",
            ["INDOOR", "OUTDOOR", "INDOOR/OUTDOOR"],
            default=st.session_state["filter_perangkat"],
            placeholder="Pilih opsi",
            key="filter_perangkat"
        )

with col3:
    filter_jenis = None
    if data_option == "Running Text":
        filter_jenis = st.multiselect(
            "🔍 Filter Jenis",
            ["LED", "RGB"],
            default=st.session_state["filter_jenis"],
            placeholder="Pilih opsi",
            key="filter_jenis"
        )
    elif data_option == "Videotron":
        filter_jenis = st.multiselect(
            "🔍 Filter Jenis",
            ["P1", "P2", "P3", "P4", "P5", "P6", "P7", "P8", "P9", "P10", "P11", "P12"],
            default=st.session_state["filter_jenis"],
            placeholder="Pilih opsi",
            key="filter_jenis"
        )

# Mulai dari data awal
filtered_df = df.copy()

# Filter berdasarkan input user
if filter_produk:
    filtered_df = filtered_df[
        filtered_df["NAMA PRODUK"].astype(str).str.contains(filter_produk, case=False, na=False)
    ]
if filter_perangkat and len(filter_perangkat) > 0:
    filtered_df = filtered_df[filtered_df["PERANGKAT"].isin(filter_perangkat)]
if filter_jenis and len(filter_jenis) > 0:
    filtered_df = filtered_df[filtered_df["JENIS"].isin(filter_jenis)]

# --- Convert kolom HARGA & TERJUAL ke numerik ---
for col in ["HARGA", "TERJUAL"]:
    if col in filtered_df.columns:
        filtered_df[col] = (
            filtered_df[col].astype(str).str.replace(r"[^\d]", "", regex=True)
        )
        filtered_df[col] = pd.to_numeric(filtered_df[col], errors="coerce")

st.subheader("Hasil Filter")
st.dataframe(filtered_df, use_container_width=True)

# --- 3. FILTER KEDUA ---
order_option_label = st.selectbox(
    "🔍 Pilih urutan",
    [ORDER_PLACEHOLDER, "Tertinggi", "Terendah"],
    key="order_option_label"
)
order_option = None if order_option_label == ORDER_PLACEHOLDER else order_option_label

# --- 4. FILTER KETIGA ---
value_option_label = st.selectbox(
    "🔍 Pilih kategori",
    [VALUE_PLACEHOLDER, "HARGA", "TERJUAL"],
    key="value_option_label"
)
value_option = None if value_option_label == VALUE_PLACEHOLDER else value_option_label

# --- 4a. DEFAULT KOLUM ---
if value_option is None and order_option is not None:
    default_col_option = st.selectbox(
        "📌 Pilih kolom default untuk urutan",
        ["TERJUAL", "HARGA"],
        key="default_col_option"
    )
else:
    default_col_option = st.session_state.get("default_col_option", "TERJUAL")

# --- FILTER AKHIR ---
final_df = filtered_df.copy()
if value_option is not None and order_option is not None:
    ascending = True if order_option == "Terendah" else False
    final_df = final_df.sort_values(by=value_option, ascending=ascending).head(10)
elif value_option is None and order_option is not None:
    ascending = True if order_option == "Terendah" else False
    final_df = final_df.sort_values(by=default_col_option, ascending=ascending).head(10)

st.subheader("Hasil Filter Akhir")
st.dataframe(final_df, use_container_width=True)

# --- RATA-RATA HARGA ---
rata_rata_harga = None
note = ""
if not final_df.empty and value_option is not None:
    rata_rata_harga = final_df["HARGA"].mean()
    if value_option == "HARGA":
        note = f"(10 {order_option})" if order_option else "(semua data)"
    elif value_option == "TERJUAL":
        note = f"(10 produk dengan {value_option} {order_option})"
    
    st.markdown(
        f"""
        <div style="background-color:#262626;padding:15px;border-radius:10px;text-align:center;
        box-shadow:2px 2px 6px rgba(0,0,0,0.5);margin-top:10px;margin-bottom:25px;">
            <h4 style="color:#FFD700;margin:0;">Rata-rata Harga Produk {note}:</h4>
            <p style="color:white;margin:0;font-size:18px;"><b>Rp {rata_rata_harga:,.0f}</b></p>
        </div>
        """,
        unsafe_allow_html=True
    )

# --- Summary Card ---
total_toko = filtered_df["TOKO"].nunique() if "TOKO" in filtered_df.columns else 0
total_produk = len(filtered_df)

if total_produk > 0:
    colA, colB = st.columns(2)
    with colA:
        st.markdown(
            f"""
            <div style="background-color:#262626;padding:20px;border-radius:15px;text-align:center;
                box-shadow:2px 2px 8px rgba(0,0,0,0.5);margin-bottom:30px;">
                <h3 style="color:#00FFAA;margin:0;">{total_toko}</h3>
                <p style="color:white;margin:0;">Total Toko</p>
            </div>
            """, unsafe_allow_html=True
        )
    with colB:
        st.markdown(
            f"""
            <div style="background-color:#262626;padding:20px;border-radius:15px;text-align:center;
                box-shadow:2px 2px 8px rgba(0,0,0,0.5);margin-bottom:30px;">
                <h3 style="color:#00FFAA;margin:0;">{total_produk}</h3>
                <p style="color:white;margin:0;">Total Produk</p>
            </div>
            """, unsafe_allow_html=True
        )

# --- WRAP TEXT ---
def wrap_text(x, width=25, max_lines=2):
    text = str(x)
    wrapped = textwrap.wrap(text, width=width)
    if len(wrapped) > max_lines:
        wrapped = wrapped[:max_lines]
        wrapped[-1] = wrapped[-1][:max(0, width-3)] + "..."
    return "<br>".join(wrapped)

final_df["NAMA PRODUK"] = final_df["NAMA PRODUK"].apply(wrap_text)

# --- GRAFIK ---
fig = None
if value_option is not None:
    y_col = value_option
elif value_option is None and order_option is not None:
    y_col = default_col_option
else:
    y_col = None

if y_col and not final_df.empty:
    fig = px.bar(
        final_df,
        x="NAMA PRODUK",
        y=y_col,
        color=y_col,
        color_continuous_scale="Viridis",
        title=f"{order_option or 'Urutan'} berdasarkan {y_col}"
              f"<br><sup>(Analisis dari {total_toko} toko dengan {total_produk} produk di Tokopedia)</sup>",
    )
    fig.update_layout(
        plot_bgcolor="black",
        paper_bgcolor="black",
        font=dict(size=10, color="white"),
        title_font=dict(size=18, color="white"),
        xaxis=dict(tickangle=-45, automargin=True),
        yaxis=dict(automargin=True)
    )
    st.plotly_chart(fig, use_container_width=True)

    # --- Download PNG ---
    buf = io.BytesIO()
    fig.write_image(buf, format="png")
    st.download_button("📥 Download Grafik (PNG)", buf.getvalue(), "grafik_tokopedia.png", "image/png")

    # --- Download HTML ---
    html_buf = io.StringIO()
    fig.write_html(html_buf)
    st.download_button("📥 Download Grafik (HTML Interaktif)", html_buf.getvalue(), "grafik_tokopedia.html", "text/html")

# --- Fungsi buat DOCX ---
def create_docx(dataframe, fig, total_toko, total_produk, value_option, order_option, rata_rata_harga=None, note=""):
    doc = Document()
    doc.add_heading('Laporan Hasil Analisis Produk Tokopedia', 0)

    # Ringkasan
    doc.add_heading("📌 Ringkasan Hasil Analisis (Produk yang dianalisis)", level=1)
    doc.add_paragraph(f"• Jumlah Toko   : {total_toko}")
    doc.add_paragraph(f"• Jumlah Produk : {total_produk}")
    if rata_rata_harga is not None:
        doc.add_paragraph(f"• Rata-rata Harga Produk {note}: Rp {rata_rata_harga:,.0f}")

    # Pilih 1 produk sesuai filter user
    if not dataframe.empty and value_option is not None and order_option is not None:
        if order_option == "Tertinggi":
            selected_row = dataframe.loc[dataframe[value_option].idxmax()]
        else:
            selected_row = dataframe.loc[dataframe[value_option].idxmin()]

        doc.add_heading(f"📌 Detail Produk dengan {value_option} {order_option}", level=1)
        doc.add_paragraph(f"• Nama Produk : {selected_row['NAMA PRODUK']}")
        if "PERANGKAT" in dataframe.columns:
            doc.add_paragraph(f"- Perangkat   : {selected_row['PERANGKAT']}")
        if "UKURAN" in dataframe.columns:
            doc.add_paragraph(f"- Ukuran      : {selected_row['UKURAN']}")
        if "JENIS" in dataframe.columns:
            doc.add_paragraph(f"- Jenis Produk: {selected_row['JENIS']}")
        if "HARGA" in dataframe.columns:
            doc.add_paragraph(f"- Harga Produk: Rp {selected_row['HARGA']:,.0f}")
        if "TERJUAL" in dataframe.columns:
            doc.add_paragraph(f"- Terjual     : {selected_row['TERJUAL']} unit")

    # --- Tambahkan Tabel Hasil Filter ---
    if not dataframe.empty:
        doc.add_heading("📊 Tabel 10 Hasil Filter Produk", level=1)
        top10 = dataframe.head(10)

        table = doc.add_table(rows=1, cols=len(top10.columns))
        table.style = "Light Grid"

        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(top10.columns):
            hdr_cells[i].text = str(col_name)

        for _, row in top10.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

    # --- Grafik hanya kalau ada fig ---
    if fig is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
            fig.write_image(tmpfile.name, format="png")
            doc.add_heading("📊 Visualisasi Grafik", level=1)
            doc.add_picture(tmpfile.name, width=Inches(6))
    else:
        doc.add_paragraph("⚠️ Grafik tidak tersedia karena filter belum dipilih atau data kosong.")

    # Simpan DOCX
    tmp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_path.name)
    return tmp_path.name

# --- Tombol Download DOCX ---
if final_df.empty:
    st.warning("⚠️ Data kosong, tidak bisa membuat laporan.")
else:
    docx_path = create_docx(final_df, fig, total_toko, total_produk, value_option, order_option, rata_rata_harga, note)
    with open(docx_path, "rb") as f:
        st.download_button(
            label="📥 Download Laporan (DOCX)",
            data=f,
            file_name="laporan_tokopedia.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
